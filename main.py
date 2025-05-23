import re
import matplotlib.pyplot as plt
import numpy as np
import tkinter as tk
from tkinter import messagebox
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Pt
from cvss_metrics import metrics, descriptions, ukrainian_labels, metric_patterns
from datetime import datetime
import os
import docx
def parse_cvss_vector(vector):
    """
    Парсинг CVSS вектора версії 3.1.

    :param vector: CVSS вектор у форматі рядка
    :return: Словник з параметрами вектора
    """
    if not vector.startswith("CVSS:3.1/"):
        raise ValueError("Invalid CVSS version, only CVSS:3.1 is supported")

    vector = vector[len("CVSS:3.1/"):]
    parts = vector.split('/')
    all_metrics = {}
    for part in parts:
        for metric, pattern in metric_patterns.items():
            match = re.match(pattern, part)
            if match:
                all_metrics[metric] = metrics[metric][match.group(1)]
                break
    return all_metrics

def get_metric_symbol_and_description(metric, value):
    """
    Визначає символ для заданої метрики і значення, та повертає український опис.

    :param metric: Метрика (наприклад, 'AV')
    :param value: Значення метрики (наприклад, 0.85)
    :return: Символ метрики та опис українською мовою
    """
    for symbol, val in metrics.get(metric, {}).items():
        if val == value:
            description = descriptions.get(metric, {}).get(symbol, '')
            return description
    return 'Невідомий символ'
def create_directory_if_not_exists(directory):
    if not os.path.exists(directory):
        os.makedirs(directory)
        print(f"Directory '{directory}' created.")
    else:
        print(f"Directory '{directory}' already exists.")
def plot_cvss_metrics_polar(metrics, frame):
    """
    Побудова полярної діаграми для метрик CVSS з підписами українською.

    :param metrics: Словник з метриками CVSS
    :param frame: Tkinter frame для відображення графіка
    """
    labels = list(metrics.keys())
    values = list(metrics.values())

    values += values[:1]
    angles = np.linspace(0, 2 * np.pi, len(labels) + 1, endpoint=True)

    fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color='blue', alpha=0.25)
    ax.plot(angles, values, color='blue', linewidth=2)

    ax.set_yticks([0.2, 0.4, 0.6, 0.8, 1.0])
    ax.set_yticklabels(['0.2', '0.4', '0.6', '0.8', '1.0'])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontdict={'fontsize': 36, 'fontname': 'Times New Roman'})

    #plt.title('Полярна діаграма CVSS метрик', size=15, color='blue', y=1.1)

    # =PNG
    current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
    to_file = re.sub(r'[^A-Za-z]', '', vector_entry.get())
    create_directory_if_not_exists("image")
    image_path = f'image/CVSS_Metrics_{current_time}_{to_file}.png'
    create_directory_if_not_exists('data')
    plt.savefig(image_path, bbox_inches='tight')
    plt.close(fig)

    canvas = FigureCanvasTkAgg(fig, master=frame)
    canvas.draw()
    canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    return image_path



def create_word_table(metrics, vector, image_path):
    """
    Створення Word документа з таблицею метрик CVSS і вставленням графіка.

    :param metrics: Словник з метриками CVSS
    :param vector: CVSS вектор у форматі рядка
    :param image_path: Шлях до зображення графіка
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    doc.add_heading('CVSS Метрики', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Код'
    hdr_cells[1].text = 'Назва'
    hdr_cells[2].text = 'Значення'
    hdr_cells[3].text = 'Опис'
    for metric, value in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = ukrainian_labels.get(metric, '')
        row_cells[2].text = str(value)
        row_cells[3].text = get_metric_symbol_and_description(metric, value)
    doc.add_heading('Полярна діаграма', level=1)
    doc.add_picture(image_path, width=docx.shared.Inches(6))

    current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
    to_file = re.sub(r'[^A-Za-z]', '', vector)
    doc.save(f'data/CVSS_Metrics__{current_time}__{to_file}.docx')

def on_submit():
    """
    Обробка події натискання кнопки 'Submit'.
    """
    vector = vector_entry.get()
    try:
        metrics = parse_cvss_vector(vector)
        for widget in chart_frame.winfo_children():
            widget.destroy()
        image_path = plot_cvss_metrics_polar(metrics, chart_frame)
        create_word_table(metrics, vector, image_path)
        messagebox.showinfo("Успіх", "CVSS вектор успішно розпарсений і Word документ створено!")
    except ValueError as e:
        messagebox.showerror("Помилка", str(e))

root = tk.Tk()
root.title("CVSS 3.1 Полярна діаграма @Anton Tykhoplav")

main_frame = tk.Frame(root)
main_frame.pack(padx=10, pady=10)

vector_label = tk.Label(main_frame, text="Введіть CVSS Вектор:")
vector_label.pack(side=tk.LEFT)

vector_entry = tk.Entry(main_frame, width=50)
vector_entry.pack(side=tk.LEFT, padx=5)

submit_button = tk.Button(main_frame, text="Побудувати", command=on_submit)
submit_button.pack(side=tk.LEFT, padx=5)

chart_frame = tk.Frame(root)
chart_frame.pack(padx=10, pady=10)

root.mainloop()
